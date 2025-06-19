import os
import re
import traceback
import uuid
import zipfile
from urllib.parse import urlparse, urljoin, unquote
import warnings
import docx

from bs4 import BeautifulSoup, Tag, NavigableString, XMLParsedAsHTMLWarning
from docx import Document
from docx.shared import Pt, Inches
from lxml import etree
from pathlib import Path

from transgemini.core.utils import get_image_extension_from_data, convert_emf_to_png, create_image_placeholder, find_image_placeholders
from transgemini.config import IMAGE_PLACEHOLDER_PREFIX, DOCX_AVAILABLE, BS4_AVAILABLE, LXML_AVAILABLE, EBOOKLIB_AVAILABLE, PILLOW_AVAILABLE


warnings.filterwarnings("ignore", category=XMLParsedAsHTMLWarning)


def read_docx_with_images(filepath, temp_dir, image_map):
    """Reads DOCX, extracts text, replaces images with placeholders, saves images."""
    if not DOCX_AVAILABLE: raise ImportError("python-docx library is required.")
    if not os.path.exists(filepath): raise FileNotFoundError(f"DOCX file not found: {filepath}")

    doc = docx.Document(filepath)
    output_lines = []

    is_bold_chapter = re.compile(r'^\s*(Глава|Chapter|Part)\s+([0-9IVXLCDM]+|[a-zA-Zа-яА-Я]+)\b.*', re.IGNORECASE)
    doc_rels = doc.part.rels
    processed_image_rids = set()
    processed_rid_to_uuid = {}

    for element in doc.element.body:

        if element.tag.endswith('p'):
            para = docx.text.paragraph.Paragraph(element, doc)
            para_text_parts = []
            contains_image = False
            for run in para.runs:

                drawing_elems = run.element.xpath('.//w:drawing')
                if drawing_elems:
                    for drawing in drawing_elems:

                        inline_elems = drawing.xpath('.//wp:inline | .//wp:anchor')
                        if inline_elems:
                            for inline in inline_elems:

                                blip_fill = inline.xpath('.//a:blip')
                                if blip_fill:
                                    rId = blip_fill[0].get(
                                        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')

                                    if rId and rId in doc_rels and "image" in doc_rels[rId].target_ref:
                                        if rId not in processed_image_rids:
                                            try:
                                                img_part = doc_rels[rId].target_part
                                                img_data = img_part.blob
                                                original_filename = os.path.basename(img_part.partname)

                                                img_ext_original = os.path.splitext(original_filename)[
                                                    -1].lower().strip('.')
                                                img_ext_detected = get_image_extension_from_data(img_data,
                                                                                                 fallback_ext=img_ext_original or "png")

                                                if img_ext_original == 'emf' or img_ext_detected == 'emf':
                                                    png_data = convert_emf_to_png(img_data)
                                                    if png_data:
                                                        img_data = png_data;
                                                        img_ext_final = 'png';
                                                        content_type = 'image/png'
                                                        print(
                                                            f"[INFO] DOCX: Converted EMF image '{original_filename}' to PNG.")
                                                    else:
                                                        print(
                                                            f"[WARN] DOCX: Failed to convert EMF '{original_filename}', skipping.");
                                                        continue  # Skip if conversion failed
                                                else:
                                                    img_ext_final = img_ext_detected;
                                                    content_type = f"image/{img_ext_final}"

                                                width, height = None, None
                                                try:
                                                    extent = inline.xpath('.//wp:extent');
                                                    if extent:
                                                        emu_per_px = 9525  # Approx conversion factor
                                                        width = int(extent[0].get('cx')) // emu_per_px
                                                        height = int(extent[0].get('cy')) // emu_per_px
                                                except Exception:
                                                    pass  # Ignore errors getting dimensions

                                                img_uuid = uuid.uuid4().hex
                                                saved_filename = f"{img_uuid}.{img_ext_final}";
                                                saved_path = os.path.join(temp_dir, saved_filename)
                                                with open(saved_path, 'wb') as img_file:
                                                    img_file.write(img_data)
                                                image_map[img_uuid] = {'saved_path': saved_path,
                                                                       'original_filename': original_filename,
                                                                       'content_type': content_type, 'width': width,
                                                                       'height': height}
                                                processed_image_rids.add(rId);
                                                processed_rid_to_uuid[rId] = img_uuid

                                                placeholder = create_image_placeholder(img_uuid)
                                                para_text_parts.append(placeholder);
                                                contains_image = True
                                            except Exception as e:
                                                print(f"[WARN] DOCX: Error processing image rId {rId}: {e}");
                                                para_text_parts.append(run.text)  # Append run text on error
                                        else:  # Image already processed (e.g., copy-pasted image)
                                            if rId in processed_rid_to_uuid:
                                                para_text_parts.append(
                                                    create_image_placeholder(processed_rid_to_uuid[rId]));
                                                contains_image = True
                                            else:  # Should not happen if processed correctly
                                                print(f"[WARN] DOCX: rId {rId} processed but not in UUID map.");
                                                para_text_parts.append(run.text)
                                    else:  # Not an image relationship or rId invalid
                                        para_text_parts.append(run.text)
                                else:  # No blip fill found
                                    para_text_parts.append(run.text)
                        else:  # No inline/anchor element
                            para_text_parts.append(run.text)
                else:  # No drawing element in run
                    para_text_parts.append(run.text)

            full_para_text = "".join(para_text_parts).strip()
            style_name = para.style.name.lower() if para.style and para.style.name else ''
            is_heading_style = False

            is_run_bold = all(r.bold for r in para.runs if r.text.strip())  # Check if all text runs are bold
            if style_name.startswith('heading 1') or (
                    style_name == 'normal' and is_bold_chapter.match(full_para_text) and is_run_bold):
                output_lines.append(f"# {full_para_text}");
                is_heading_style = True
            elif style_name.startswith('heading 2'):
                output_lines.append(f"## {full_para_text}");
                is_heading_style = True
            elif style_name.startswith('heading 3'):
                output_lines.append(f"### {full_para_text}");
                is_heading_style = True

            elif not full_para_text.strip() and not contains_image:

                if output_lines and output_lines[-1] != "": output_lines.append("")
                continue

            elif not is_heading_style and (style_name.startswith('list paragraph') or (
                    para.paragraph_format and para.paragraph_format.left_indent and full_para_text)):
                list_marker = "*";  # Default marker

                num_match = re.match(r'^\s*(\d+\.|\([a-z]\)|\([A-Z]\)|[a-z]\.|[A-Z]\.)\s+',
                                     para.text)  # Numbered or lettered lists
                bullet_match = re.match(r'^\s*([\*\-\•\⁃])\s+', para.text)  # Common bullet chars
                if num_match:
                    list_marker = num_match.group(1)
                elif bullet_match:
                    list_marker = bullet_match.group(1)

                clean_list_text = re.sub(r'^\s*(\d+\.|\([a-z]\)|\([A-Z]\)|[a-z]\.|[A-Z]\.|[\*\-\•\⁃])\s*', '',
                                         full_para_text)
                output_lines.append(f"{list_marker} {clean_list_text}")

            elif not is_heading_style and (full_para_text or contains_image):
                output_lines.append(full_para_text)

        elif element.tag.endswith('tbl'):

            if output_lines and output_lines[-1]: output_lines.append("")
            output_lines.append("[--- ТАБЛИЦА (не обработано) ---]")
            output_lines.append("")

    final_text = "";
    for i, line in enumerate(output_lines):
        final_text += line

        if i < len(output_lines) - 1:
            final_text += "\n"

            is_current_placeholder_line = IMAGE_PLACEHOLDER_PREFIX in line
            is_next_placeholder_line = IMAGE_PLACEHOLDER_PREFIX in output_lines[i + 1]
            is_current_heading = line.startswith('#')
            is_current_list = re.match(r'^([\*\-\•\⁃]|\d+\.|\([a-z]\)|\([A-Z]\)|[a-z]\.|[A-Z]\.)\s', line)
            is_current_table = "[--- ТАБЛИЦА" in line
            is_next_table = "[--- ТАБЛИЦА" in output_lines[i + 1]

            if (output_lines[i + 1] != "" and line != "" and  # Both lines have content
                    not is_current_heading and not is_current_list and  # Not headings or lists
                    not is_current_table and not is_next_table and  # Not tables
                    not (is_current_placeholder_line and is_next_placeholder_line)):  # Not two image lines together
                final_text += "\n"

    print(f"[INFO] DOCX Read: Extracted {len(image_map)} images.")
    return final_text.strip()

def process_html_images(html_content, source_context, temp_dir, image_map):
    """
    Parses HTML, extracts images, replaces with placeholders, converts Hx/title to Markdown-like,
    and then extracts text content for translation.
    `source_context` can be a tuple (zipfile.ZipFile, html_path_in_zip) or a base directory path.
    """
    if not BS4_AVAILABLE: raise ImportError("BeautifulSoup4 is required for HTML processing.")

    if "<svg" in html_content.lower() or "xmlns:" in html_content.lower() or \
            html_content.strip().startswith("<?xml"):
        parser_type = 'lxml-xml'  # Use 'lxml-xml' for stricter XML or documents with SVG/namespaces
    else:
        parser_type = 'lxml'  # Use 'lxml' for general HTML

    try:
        soup = BeautifulSoup(html_content, parser_type)
    except Exception as e_parse:
        print(f"DEBUG process_html_images: Parse failed with '{parser_type}': {e_parse}. Trying 'html.parser'.")
        try:
            soup = BeautifulSoup(html_content, 'html.parser')  # Fallback parser
        except Exception as e_parse_fallback:
            print(
                f"[ERROR] BeautifulSoup failed to parse HTML content with primary parser '{parser_type}' and fallback 'html.parser'. Error: {e_parse_fallback}")
            raise ValueError(f"Failed to parse HTML content after trying multiple parsers: {e_parse_fallback}")

    zip_file_obj = None
    source_html_path = None  # e.g., OEBPS/Text/0005_SE1000.xhtml
    base_path = ""  # e.g., OEBPS
    if isinstance(source_context, tuple) and len(source_context) == 2 and isinstance(source_context[0],
                                                                                     zipfile.ZipFile):
        zip_file_obj = source_context[0]
        source_html_path = source_context[1]
        if source_html_path:
            base_path = os.path.dirname(source_html_path).replace('\\', '/')
            if base_path == '.': base_path = ""  # Handle root case
    elif isinstance(source_context, str) and os.path.isdir(source_context):
        base_path = source_context  # This is the HTML's dir path.
        source_html_path = "unknown.html"  # Assign dummy path if only dir context given
        zip_file_obj = None
    else:
        zip_file_obj = None
        base_path = ""  # HTML dir path unknown
        source_html_path = "unknown.html"

    image_processing_context = zip_file_obj if zip_file_obj else base_path
    images_found_and_processed = 0  # Счетчик для лога

    potential_image_tags = soup.find_all(['img', 'svg'])

    for tag_index, tag in enumerate(potential_image_tags):

        if not tag.parent:  # Tag might have been removed or replaced in a previous iteration

            continue

        img_uuid = None
        tag_name = tag.name.lower()
        placeholder_str = ""

        try:
            if tag_name == 'img':

                img_uuid = _process_single_image(tag, image_processing_context, base_path, source_html_path, temp_dir,
                                                 image_map, is_svg_image=False)
                if img_uuid:
                    placeholder_str = create_image_placeholder(img_uuid)

                    tag.replace_with(NavigableString(placeholder_str))  # Use NavigableString for replacement
                    images_found_and_processed += 1
                else:

                    tag.replace_with("")  # Use empty string for removal

            elif tag_name == 'svg':

                svg_image_tag = tag.find('image', recursive=False)  # Try direct find first
                if not svg_image_tag:  # If not found, try case-insensitive search
                    svg_image_tag = tag.find(lambda t: t.name.lower() == 'image', recursive=False)

                if svg_image_tag:

                    img_uuid = _process_single_image(svg_image_tag, image_processing_context, base_path,
                                                     source_html_path, temp_dir, image_map, is_svg_image=True)
                    if img_uuid:
                        placeholder_str = create_image_placeholder(img_uuid)

                        tag.replace_with(NavigableString(placeholder_str))  # Use NavigableString
                        images_found_and_processed += 1
                    else:

                        tag.replace_with("")
                else:  # No <image> found in <svg>

                    tag.replace_with("")
        except Exception as replace_err:
            print(f"[ERROR] process_html_images: Error replacing tag <{tag_name}>: {replace_err}. Attempting removal.")
            traceback.print_exc()  # Print stack trace for replacement error
            try:
                if tag.parent:  # Check again if parent exists before trying to remove
                    tag.replace_with("")
            except Exception as remove_err:
                print(f"[ERROR] process_html_images: Failed to remove tag <{tag_name}> after error: {remove_err}")

    html_doctitle_text = None

    if soup.head and soup.head.title and soup.head.title.string:
        title_candidate = soup.head.title.string.strip()

        generic_titles_lower = [
            'untitled', 'unknown', 'navigation', 'toc', 'table of contents', 'index',
            'contents', 'оглавление', 'содержание', 'индекс',
            'cover', 'title page', 'copyright', 'chapter'  # Added more generic terms
        ]

        if title_candidate and title_candidate.lower() not in generic_titles_lower and len(
                title_candidate) > 2:  # Min length
            html_doctitle_text = title_candidate

    content_extraction_root = soup.body if soup.body else soup
    if not content_extraction_root:
        print("[WARN] process_html_images: No <body> or root element found in parsed soup after image processing.")
        return ""  # Or raise error, or return placeholder text

    for level in range(6, 0, -1):  # Process H6 down to H1 to correctly handle if headers are nested (though unusual)
        for header_tag in content_extraction_root.find_all(f'h{level}'):

            header_text_content = header_tag.get_text(separator=' ', strip=True)
            if header_text_content:
                md_prefix = NavigableString(f"\n{'#' * level} ")
                md_suffix = NavigableString("\n")

                header_tag.insert_before(md_prefix)
                header_tag.insert_after(md_suffix)

                header_tag.unwrap()

    tags_to_decompose_finally = ['script', 'style', 'noscript', 'head', 'meta', 'link', 'applet', 'embed', 'object',
                                 'form', 'iframe', 'map', 'area', 'header', 'footer', 'nav', 'aside', 'figure',
                                 'figcaption']

    for tag_type in tags_to_decompose_finally:
        decomposed_count = 0
        for instance in content_extraction_root.find_all(tag_type):
            instance.decompose()
            decomposed_count += 1

    body_text_md = content_extraction_root.get_text(separator='\n', strip=True)

    final_text_for_api = body_text_md
    if html_doctitle_text:

        body_starts_with_title_as_h1 = False
        if body_text_md.lstrip().startswith("# "):  # Check if it starts with any H1
            first_line_of_body = body_text_md.lstrip().split('\n', 1)[0]

            if first_line_of_body[2:].strip().lower() == html_doctitle_text.lower():
                body_starts_with_title_as_h1 = True

        if not body_starts_with_title_as_h1:
            final_text_for_api = f"{'#'} {html_doctitle_text}\n\n{body_text_md}"

    final_text_for_api = re.sub(r'\n{3,}', '\n\n', final_text_for_api).strip()

    final_placeholders_for_api = find_image_placeholders(final_text_for_api)

    return final_text_for_api

def _process_single_image(img_tag, source_context, base_path, source_html_path, temp_dir, image_map,
                          is_svg_image=False):
    """
    Processes individual image tag.
    For EPUB->EPUB: Extracts original src and attributes, stores them in image_map with a UUID. Does NOT save file.
    For other modes: Extracts image data, saves to temp_dir, stores path and info in image_map.
    """
    src = None
    xlink_namespace_uri = "http://www.w3.org/1999/xlink"

    is_epub_rebuild_mode = isinstance(source_context, zipfile.ZipFile)  # True if processing for EPUB->EPUB

    if is_svg_image:
        src = img_tag.get(f'{{{xlink_namespace_uri}}}href')
        if not src:
            attrs_dict = img_tag.attrs
            if 'xlink:href' in attrs_dict:
                src = attrs_dict['xlink:href']
            elif 'href' in attrs_dict:
                src = attrs_dict['href']
            else:
                namespaced_key = (xlink_namespace_uri, 'href')
                if namespaced_key in attrs_dict: src = attrs_dict[namespaced_key]

    else:  # HTML <img> tag
        src = img_tag.get('src', '')

    if not src or src.startswith('data:'):
        return None

    img_uuid = uuid.uuid4().hex
    original_src_value = src  # This is the raw value from the attribute, e.g., "../Images/0004.png"
    original_tag_name = img_tag.name  # 'img' or 'image' (from svg)
    all_original_attributes = dict(img_tag.attrs)  # Store all attributes

    if is_epub_rebuild_mode:

        image_map[img_uuid] = {
            'original_src': original_src_value,
            'original_tag_name': original_tag_name,  # 'img' or 'image'
            'is_svg_image_child': is_svg_image,  # True if it was <image> inside <svg>
            'attributes': all_original_attributes  # Store all original attributes
        }

        return img_uuid
    else:

        img_data = None
        decoded_src = unquote(src)
        original_filename = os.path.basename(urlparse(decoded_src).path)
        if not original_filename:
            src_parts = decoded_src.split('/')
            potential_fname = src_parts[-1] if src_parts else "image"
            potential_fname = potential_fname.split('?')[0]
            safe_fname_part = re.sub(r'[^\w\.\-]+', '_', potential_fname)
            _, ext_guess = os.path.splitext(safe_fname_part)
            fallback_ext = "png"
            if ext_guess and ext_guess[1:].lower() in ['jpg', 'jpeg', 'gif', 'webp', 'bmp', 'tiff', 'svg']:
                original_filename = safe_fname_part
            else:
                original_filename = f"{Path(safe_fname_part).stem}.{fallback_ext}"

        content_type = None

        try:
            if isinstance(source_context, zipfile.ZipFile):  # EPUB source -> non-EPUB output

                possible_paths = []
                current_html_dir = base_path
                path1 = os.path.join(current_html_dir, decoded_src)
                path1_norm = os.path.normpath(path1).replace('\\', '/')
                if not path1_norm.startswith('..'): possible_paths.append(path1_norm.lstrip('/'))
                path2_norm = os.path.normpath(decoded_src.lstrip('/')).replace('\\', '/')
                if not path2_norm.startswith('..'): possible_paths.append(path2_norm)

                unique_paths = list(dict.fromkeys(p.strip('/') for p in possible_paths if p.strip('/')))

                for img_path_in_zip in unique_paths:
                    try:
                        img_data = source_context.read(img_path_in_zip)

                        break
                    except KeyError:
                        continue
                if not img_data:
                    print(f"[WARN] HTML Image NOT Found (EPUB src -> File Output): src='{src}'. Tried: {unique_paths}.")
                    return None

            elif isinstance(source_context, str) and os.path.isdir(source_context):  # Directory context

                paths_to_try_fs = [
                    os.path.normpath(os.path.join(base_path, decoded_src)),
                    os.path.normpath(os.path.join(source_context, decoded_src.lstrip('/\\')))
                ]
                if not decoded_src.startswith(('/', '\\')):
                    path3_fs = os.path.normpath(os.path.join(source_context, decoded_src))
                    if path3_fs not in paths_to_try_fs: paths_to_try_fs.append(path3_fs)
                abs_path = next((p for p in paths_to_try_fs if os.path.exists(p)), None)
                if not abs_path:
                    print(f"[WARN] HTML Image (FS Mode): Could not find '{decoded_src}'. Tried: {paths_to_try_fs}")
                    return None
                with open(abs_path, 'rb') as f:
                    img_data = f.read()
            else:
                print(f"[WARN] HTML Image: Unknown source context for file mode: {type(source_context)}")
                return None

            img_ext_from_file = os.path.splitext(original_filename)[1][1:].lower()
            content_type = f"image/{get_image_extension_from_data(img_data, fallback_ext=img_ext_from_file or 'jpeg')}"
            img_ext = content_type.split('/')[-1] if content_type else 'jpeg'
            img_ext = 'jpg' if img_ext == 'jpeg' else img_ext

            if img_ext == 'emf':
                converted_data = convert_emf_to_png(img_data)
                if converted_data:
                    img_data = converted_data;
                    img_ext = 'png';
                    content_type = 'image/png'
                else:
                    return None

            filename = f"{img_uuid}.{img_ext}"
            save_path = os.path.join(temp_dir, filename)
            with open(save_path, 'wb') as f:
                f.write(img_data)

            image_map[img_uuid] = {
                'saved_path': save_path,  # For non-EPUB rebuild, this is used
                'original_filename': original_filename,
                'original_src': original_src_value,  # Still store original_src for consistency if needed
                'content_type': content_type,
                'attributes': all_original_attributes  # Store original attributes
            }

            return img_uuid

        except Exception as e:
            print(f"[ERROR] HTML Image (File Mode): Error processing src '{src}': {e}")
            traceback.print_exc()
            return None

def write_markdown_to_docx(filepath, md_text_with_placeholders, image_map):
    """Writes Markdown-like text with placeholders back to DOCX."""
    if not DOCX_AVAILABLE: raise ImportError("python-docx library is required.")
    if image_map is None: image_map = {}
    doc = Document()

    lines = re.split('(\n)', md_text_with_placeholders)

    paragraphs_md = []
    current_para_lines = []
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:  # Treat empty line as paragraph break
            if current_para_lines:
                paragraphs_md.append("\n".join(current_para_lines))
                current_para_lines = []

            if paragraphs_md and paragraphs_md[-1]:  # Add if last added wasn't already empty
                paragraphs_md.append("")
        else:
            current_para_lines.append(line)
    if current_para_lines:  # Add last paragraph if exists
        paragraphs_md.append("\n".join(current_para_lines))

    current_docx_para = None
    for md_para in paragraphs_md:
        md_para_stripped = md_para.strip()
        if not md_para_stripped:

            if current_docx_para is not None:  # Check if previous para exists
                doc.add_paragraph("")
            current_docx_para = None  # Reset current para tracker
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.*)', md_para_stripped, re.DOTALL)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text_raw = heading_match.group(2).strip()
            current_docx_para = doc.add_heading("", level=max(1, min(level, 6)))  # Add heading
            process_text_with_placeholders(current_docx_para, heading_text_raw, image_map)
            continue  # Move to next md paragraph

        list_match = re.match(r'^([\*\-\•\⁃]|\d+\.|\([a-z]\)|\([A-Z]\)|[a-z]\.|[A-Z]\.)\s+(.*)', md_para_stripped,
                              re.DOTALL)
        if list_match:
            marker = list_match.group(1)
            list_item_text_raw = list_match.group(2).strip()
            style = 'List Bullet' if marker in ['*', '-', '•', '⁃'] else 'List Number'  # Basic style mapping
            try:
                current_docx_para = doc.add_paragraph(style=style)
            except (KeyError, ValueError):  # Fallback if style doesn't exist in template
                print(f"[WARN] DOCX Write: Style '{style}' not found. Using default paragraph.")
                current_docx_para = doc.add_paragraph()
            process_text_with_placeholders(current_docx_para, list_item_text_raw, image_map)
            continue  # Move to next md paragraph

        if md_para_stripped == '---':
            doc.add_paragraph().add_run()._element.xpath('.//w:pPr')[0].append(
                etree.fromstring(
                    '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>')
            )
            current_docx_para = None  # HR acts as break
            continue

        current_docx_para = doc.add_paragraph()
        process_text_with_placeholders(current_docx_para, md_para.strip(),
                                       image_map)  # Process original (not stripped) to keep internal newlines

    doc.save(filepath)

def process_text_with_placeholders(docx_paragraph, text_with_placeholders, image_map):
    """Adds runs of text and images to a docx paragraph based on placeholders."""

    last_index = 0
    placeholders_found = find_image_placeholders(text_with_placeholders)

    if not placeholders_found:
        if text_with_placeholders.strip():
            docx_paragraph.add_run(text_with_placeholders)
        return

    for placeholder_tag, img_uuid in placeholders_found:
        match_start = text_with_placeholders.find(placeholder_tag, last_index)
        if match_start == -1: continue  # Should not happen with finditer logic

        text_before = text_with_placeholders[last_index:match_start]
        if text_before:
            docx_paragraph.add_run(text_before)

        if img_uuid in image_map:
            img_info = image_map[img_uuid];
            img_path = img_info['saved_path']
            if os.path.exists(img_path):
                try:

                    img_width_px = img_info.get('width');
                    img_height_px = img_info.get('height')
                    run = docx_paragraph.add_run()  # Create run for the picture
                    target_width = None

                    if img_width_px:
                        try:
                            img_width_px = float(img_width_px)  # Ensure it's a number
                            if img_width_px > 0:
                                target_width_inches = img_width_px / 96.0  # Approx DPI
                                max_doc_width_inches = 6.0  # Usable width on standard page
                                target_width = Inches(min(target_width_inches, max_doc_width_inches))
                        except (ValueError, TypeError):
                            pass  # Ignore invalid width values

                    run.add_picture(img_path, width=target_width)
                except FileNotFoundError:
                    print(f"[ERROR] DOCX Write: Image file not found: {img_path}")
                    docx_paragraph.add_run(
                        f"[Image NF: {img_info.get('original_filename', img_uuid)}]")  # Add error text
                except Exception as e:
                    print(f"[ERROR] DOCX Write: Failed to add picture {img_path}: {e}")
                    docx_paragraph.add_run(f"[Img Err: {img_info.get('original_filename', img_uuid)}]")
            else:

                print(f"[ERROR] DOCX Write: Image path from map does not exist: {img_path}")
                docx_paragraph.add_run(f"[Img Path Miss: {img_info.get('original_filename', img_uuid)}]")
        else:

            print(f"[WARN] DOCX Write: Placeholder UUID '{img_uuid}' not found in image_map.")
            docx_paragraph.add_run(f"[Unk Img: {img_uuid}]")

        last_index = match_start + len(placeholder_tag)

    text_after = text_with_placeholders[last_index:]
    if text_after:
        docx_paragraph.add_run(text_after)